"""
test_real_doc.py - 실제 API를 호출하는 전체 파이프라인 테스트

Mock 없이 실제 검색 API(Tavily, Naver)와 LLM(Ollama)을 호출하여
전체 파이프라인을 검증합니다.

사용법:
    python test_real_doc.py
"""

from datetime import datetime

from pipeline.pipeline_logger import PipelineContext, log_step, print_summary

# 실제 파이프라인 모듈 임포트
from pipeline.section_planner import plan_sections
from pipeline.classifier import classify_source_type
from pipeline.rag_tool import search_internal_knowledge
from pipeline.web_search import search_web
from pipeline.llm_writer import generate_section_draft
from pipeline.router import route, Route
from pipeline.docx_builder import (
    create_document,
    build_cover_page,
    build_toc_page,
    start_body_content,
    add_section,
    add_references_section,
    collect_references,
    save_document,
    TOC_THRESHOLD,
)


def run_real_pipeline(
    user_request: str = "AWS 클라우드 비용 최적화 방안에 대한 기술 보고서를 작성해줘",
    output_path: str | None = None,
    author: str = "Grounded Report",
    verbose: bool = True,
    return_summary: bool = False,
) -> str | tuple[str, dict]:
    """
    실제 API를 호출하는 전체 파이프라인 실행

    Args:
        user_request: 사용자 요청 문자열
        output_path: 출력 파일 경로 (None이면 자동 생성)
        author: 문서 작성자
        verbose: 콘솔 출력 여부
        return_summary: True면 (output_path, summary_dict) 반환

    Returns:
        str: 생성된 문서 파일 경로 (return_summary=False)
        tuple[str, dict]: (파일 경로, 실행 요약) (return_summary=True)
    """
    import hashlib
    import time

    # output_path 자동 생성
    if output_path is None:
        timestamp = int(time.time())
        hash_suffix = hashlib.md5(user_request.encode()).hexdigest()[:6]
        output_path = f"output/report_{timestamp}_{hash_suffix}.docx"

    # 파이프라인 컨텍스트 (JSON 스트리밍 끔 - 깔끔한 출력)
    ctx = PipelineContext("real_pipeline_test", stream_to_console=verbose)

    # 각 단계를 데코레이터로 래핑
    _plan_sections = log_step(ctx, "1. 섹션 계획 생성 (LLM)")(plan_sections)
    _classify = log_step(ctx, "2. 소스 유형 분류 (LLM)")(classify_source_type)
    _search_internal = log_step(ctx, "3. 내부 RAG 검색")(search_internal_knowledge)
    _search_web = log_step(ctx, "4. 웹 검색 (Tavily+Naver)")(search_web)
    _generate_draft = log_step(ctx, "5. LLM 초안 생성")(generate_section_draft)
    _route = log_step(ctx, "6. 라우팅 결정")(route)

    if verbose:
        print("=" * 70)
        print("  실제 파이프라인 테스트 (API 호출 포함)")
        print("=" * 70)
        print(f"  요청: {user_request}")
        print(f"  출력: {output_path}")
        print("=" * 70 + "\n")

    try:
        # 1. 섹션 계획 생성 (문서 제목 포함)
        sections, document_title, detected_hint = _plan_sections(user_request)

        # 문서 제목이 없으면 fallback
        if not document_title:
            document_title = "AWS 클라우드 비용 최적화 방안"

        if verbose:
            print(f"[INFO] 문서 제목: {document_title}")
            print(f"[INFO] 생성된 섹션: {[s['title'] for s in sections]}")
            print(f"[INFO] 감지된 문서 유형: {detected_hint}\n")

        # 2. 문서 초기화 (새로운 페이지 구조)
        doc = create_document()  # 빈 문서 생성 (스타일만 적용)

        # 페이지 1: 표지
        build_cover_page(
            doc,
            title=document_title,
            author=author,
            date=datetime.now().strftime("%Y-%m-%d"),
        )

        # 페이지 2: 목차 (섹션 수가 충분하면)
        include_toc = len(sections) >= TOC_THRESHOLD
        if include_toc:
            build_toc_page(doc, sections=sections)

        # 페이지 3+: 본문 시작
        start_body_content(doc)

        # 3. 각 섹션 처리 (출처 수집)
        all_sources = []
        section_summaries = []  # 결론 작성용 앞 섹션 요약

        for i, section in enumerate(sections):
            title = section["title"]
            query = section["query"]
            if verbose:
                print(f"\n[처리 중] 섹션 {i+1}/{len(sections)}: {title}")

            # 마지막 섹션만 결론으로 처리 (source_type="none", 앞 섹션 요약 주입)
            # 중간 섹션은 키워드와 관계없이 classifier 결과를 따름
            is_last_section = (i == len(sections) - 1)

            if is_last_section:
                source_type = "none"
                if verbose:
                    print(f"  → 소스 유형: none (마지막 섹션 = 결론)")
            else:
                source_type = _classify(
                    section_title=title,
                    section_query=query,
                )
                if verbose:
                    print(f"  → 소스 유형: {source_type}")

            # 분류 결과에 따라 검색
            # 쿼리에 topic 키워드가 없으면 추가 (검색 범위 확장 방지)
            search_query = query
            if document_title and document_title not in query:
                # topic에서 핵심 키워드 추출하여 쿼리 앞에 추가
                search_query = f"{document_title} - {query}"

            sources = []
            if source_type == "internal":
                result = _search_internal(search_query)
                sources = result.get("results", [])
                if verbose:
                    print(f"  → 내부 검색 결과: {len(sources)}건")
            elif source_type == "web":
                result = _search_web(search_query, topic=document_title)
                sources = result.get("results", [])
                if verbose:
                    print(f"  → 웹 검색 결과: {len(sources)}건")
            else:
                if verbose:
                    print(f"  → 검색 불필요 (source_type=none)")

            # 출처 수집 (나중에 참고 자료 섹션에 추가)
            if sources:
                all_sources.append(sources)

            # LLM 초안 생성 (마지막 섹션이면 앞 섹션 요약 전달)
            prev_summary = "\n\n".join(section_summaries) if is_last_section else None

            llm_result = _generate_draft(
                section_query=query,
                sources=sources,
                source_type=source_type,
                topic=document_title,
                previous_sections_summary=prev_summary,
            )
            if verbose:
                print(f"  → LLM 생성 완료 (답변 길이: {len(llm_result.get('answer', ''))}자)")

            # 앞 섹션 본문 수집 (결론 제외) - 각 섹션 앞 300자
            if not is_last_section and llm_result.get("answer"):
                answer = llm_result["answer"]
                # 앞 2~3문장 (약 300자) 추출
                excerpt = answer[:300].rsplit(".", 1)[0] + "." if len(answer) > 300 else answer
                section_summaries.append(f"[{title}]\n{excerpt}")

            # 라우팅 결정
            decision = _route(
                source_type=llm_result["source_type"],
                source_count=llm_result["source_count"],
                source_relevance=llm_result["source_relevance"],
                has_fabrication_risk=llm_result["has_fabrication_risk"],
            )
            if verbose:
                print(f"  → 라우팅: {decision.label}")

            # 섹션 추가 (인라인 출처 표시 안 함)
            add_section(doc, title, llm_result, decision, sources, include_inline_sources=False)

        # 4. 참고 자료 섹션 추가 (모든 출처 통합)
        unique_sources = collect_references(all_sources)
        if unique_sources:
            add_references_section(doc, unique_sources)
            if verbose:
                print(f"\n[INFO] 참고 자료: {len(unique_sources)}건 (중복 제거됨)")

        # 5. 문서 저장
        save_document(doc, output_path)

        if verbose:
            print(f"\n[성공] 문서 저장 완료: {output_path}")

    except Exception as e:
        if verbose:
            print(f"\n[에러] 파이프라인 실패: {e}")
            import traceback
            traceback.print_exc()
        raise  # 예외를 다시 발생시켜 호출자가 처리할 수 있게 함

    # 타임라인 요약 출력
    if verbose:
        print_summary(ctx, detailed=False)

    # 요약 반환 옵션
    if return_summary:
        ctx.finish()
        return output_path, ctx.get_summary()

    return output_path


def verify_output(output_path: str):
    """생성된 문서 검증"""
    import os
    from docx import Document

    print("\n" + "=" * 70)
    print("  생성된 문서 검증")
    print("=" * 70)

    # 파일 존재 확인
    if not os.path.exists(output_path):
        print(f"[실패] 파일이 존재하지 않음: {output_path}")
        return False

    file_size = os.path.getsize(output_path)
    print(f"[확인] 파일 존재: {output_path} ({file_size:,} bytes)")

    # 문서 내용 확인
    doc = Document(output_path)

    # 단락 수
    para_count = len(doc.paragraphs)
    print(f"[확인] 문단 수: {para_count}")

    # 표 수
    table_count = len(doc.tables)
    print(f"[확인] 표 수: {table_count}")

    # Heading 추출
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    print(f"[확인] 섹션 제목들:")
    for h in headings[:10]:
        print(f"       - {h[:50]}{'...' if len(h) > 50 else ''}")

    # 본문 샘플 (첫 번째 일반 문단)
    body_paras = [p.text for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    if body_paras:
        sample = body_paras[0][:200]
        print(f"\n[샘플] 본문 첫 문단:")
        print(f"       {sample}...")

    print("\n[성공] 문서 검증 완료")
    return True


if __name__ == "__main__":
    output_path = run_real_pipeline()
    verify_output(output_path)
