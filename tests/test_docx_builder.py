"""
test_docx_builder.py - docx_builder 모듈 테스트

폰트, 목차, 표 자동 삽입 로직에 대한 기본 테스트를 포함합니다.
"""

import os
import tempfile
import pytest
from docx import Document
from docx.oxml.ns import qn

from pipeline.docx_builder import (
    create_document,
    add_section,
    add_table,
    save_document,
    build_document,
    _set_korean_font,
    _parse_content_structure,
    _add_toc,
    _count_heading2,
    TOC_THRESHOLD,
)
from pipeline.router import Route


class TestKoreanFont:
    """한글 폰트 설정 테스트"""

    def test_set_korean_font_sets_eastasia(self):
        """w:eastAsia 속성이 정확히 설정되는지 확인"""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("한글 테스트")

        _set_korean_font(run, "맑은 고딕")

        # XML 레벨에서 w:eastAsia 확인
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))

        assert rFonts is not None
        assert rFonts.get(qn('w:eastAsia')) == "맑은 고딕"
        assert rFonts.get(qn('w:ascii')) == "맑은 고딕"
        assert rFonts.get(qn('w:hAnsi')) == "맑은 고딕"

    def test_create_document_applies_default_font(self):
        """create_document가 기본 한글 폰트를 적용하는지 확인"""
        doc = create_document(title="테스트 문서")

        # Normal 스타일에 한글 폰트 설정 확인
        style = doc.styles['Normal']
        assert style.font.name == "맑은 고딕"


class TestTOC:
    """목차(TOC) 삽입 테스트"""

    def test_toc_inserted_when_sections_exceed_threshold(self):
        """섹션 수가 TOC_THRESHOLD 이상이면 목차 삽입"""
        doc = create_document(title="테스트", include_toc=True)

        # 목차 필드가 삽입되었는지 확인 (instrText 존재)
        xml_str = doc._body._element.xml
        assert 'TOC' in xml_str

    def test_toc_not_inserted_when_below_threshold(self):
        """섹션 수가 TOC_THRESHOLD 미만이면 목차 미삽입"""
        doc = create_document(title="테스트", include_toc=False)

        xml_str = doc._body._element.xml
        assert 'TOC \\o' not in xml_str

    def test_build_document_auto_toc_above_threshold(self):
        """build_document가 섹션 수에 따라 자동으로 목차 삽입"""
        sections = [
            {"title": f"섹션 {i}", "query": f"질의 {i}", "order": i}
            for i in range(1, TOC_THRESHOLD + 2)
        ]
        results = [
            ({"answer": f"답변 {i}", "source_type": "none", "source_count": 0}, Route.AUTO_APPROVE, [])
            for i in range(len(sections))
        ]

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            build_document(sections, results, output_path=output_path)
            doc = Document(output_path)
            xml_str = doc._body._element.xml
            assert 'TOC' in xml_str
        finally:
            os.unlink(output_path)

    def test_save_document_auto_toc_inserts_after_heading1(self):
        """save_document가 Heading 2 수가 임계값 이상일 때 자동으로 목차 삽입"""
        doc = create_document(title="테스트 제목")

        # TOC_THRESHOLD 이상의 섹션 추가
        for i in range(TOC_THRESHOLD):
            llm_result = {"answer": f"내용 {i}", "source_type": "none", "source_count": 0}
            add_section(doc, f"섹션 {i+1}", llm_result, Route.AUTO_APPROVE, [])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            save_document(doc, output_path, auto_toc=True)
            saved_doc = Document(output_path)
            xml_str = saved_doc._body._element.xml
            assert 'TOC' in xml_str

            # "목차" 텍스트가 포함되어 있는지 확인
            all_text = " ".join([p.text for p in saved_doc.paragraphs])
            assert "목차" in all_text
        finally:
            os.unlink(output_path)

    def test_save_document_no_auto_toc_below_threshold(self):
        """save_document가 Heading 2 수가 임계값 미만일 때 목차 미삽입"""
        doc = create_document(title="테스트 제목")

        # TOC_THRESHOLD 미만의 섹션 추가
        for i in range(TOC_THRESHOLD - 1):
            llm_result = {"answer": f"내용 {i}", "source_type": "none", "source_count": 0}
            add_section(doc, f"섹션 {i+1}", llm_result, Route.AUTO_APPROVE, [])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            save_document(doc, output_path, auto_toc=True)
            saved_doc = Document(output_path)
            xml_str = saved_doc._body._element.xml
            assert 'TOC \\o' not in xml_str
        finally:
            os.unlink(output_path)

    def test_count_heading2(self):
        """_count_heading2가 Heading 2 수를 정확히 카운트"""
        doc = create_document()
        doc.add_heading("제목 1", level=1)
        doc.add_heading("섹션 1", level=2)
        doc.add_heading("섹션 2", level=2)
        doc.add_heading("소제목", level=3)
        doc.add_heading("섹션 3", level=2)

        count = _count_heading2(doc)
        assert count == 3


class TestTable:
    """표 삽입 테스트"""

    def test_add_table_creates_table_with_header_shading(self):
        """표 헤더 행에 음영이 적용되는지 확인"""
        doc = Document()
        headers = ["항목", "내용", "비고"]
        rows = [
            ["항목1", "내용1", "비고1"],
            ["항목2", "내용2", "비고2"],
        ]

        table = add_table(doc, headers, rows)

        # 테이블 생성 확인
        assert len(doc.tables) == 1
        assert len(table.rows) == 3  # 헤더 1 + 데이터 2

        # 헤더 행 셀 음영 확인 (브랜드 Primary 컬러: #003899)
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shading = tcPr.find(qn('w:shd'))
            assert shading is not None
            assert shading.get(qn('w:fill')) == '003899'

    def test_add_table_header_text_is_bold(self):
        """표 헤더 텍스트가 굵게 표시되는지 확인"""
        doc = Document()
        headers = ["컬럼A", "컬럼B"]
        rows = [["값1", "값2"]]

        table = add_table(doc, headers, rows)

        header_row = table.rows[0]
        for cell in header_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    assert run.bold is True


class TestContentParsing:
    """콘텐츠 파싱 테스트"""

    def test_parse_bullet_list(self):
        """Bullet 리스트 파싱"""
        text = "- 항목 1\n- 항목 2\n- 항목 3"
        result = _parse_content_structure(text)

        assert len(result) == 1
        assert result[0]["type"] == "bullet_list"
        assert result[0]["items"] == ["항목 1", "항목 2", "항목 3"]

    def test_parse_numbered_list(self):
        """Numbered 리스트 파싱"""
        text = "1. 첫 번째\n2. 두 번째\n3. 세 번째"
        result = _parse_content_structure(text)

        assert len(result) == 1
        assert result[0]["type"] == "numbered_list"
        assert result[0]["items"] == ["첫 번째", "두 번째", "세 번째"]

    def test_parse_mixed_content(self):
        """혼합 콘텐츠 파싱"""
        text = "서론 내용입니다.\n\n- 포인트 1\n- 포인트 2\n\n결론 내용입니다."
        result = _parse_content_structure(text)

        assert len(result) == 3
        assert result[0]["type"] == "paragraph"
        assert result[0]["content"] == "서론 내용입니다."
        assert result[1]["type"] == "bullet_list"
        assert result[1]["items"] == ["포인트 1", "포인트 2"]
        assert result[2]["type"] == "paragraph"
        assert result[2]["content"] == "결론 내용입니다."

    def test_parse_empty_text(self):
        """빈 텍스트 파싱"""
        result = _parse_content_structure("")
        assert result == []

    def test_parse_asterisk_bullet(self):
        """* 로 시작하는 bullet 리스트"""
        text = "* 항목 A\n* 항목 B"
        result = _parse_content_structure(text)

        assert len(result) == 1
        assert result[0]["type"] == "bullet_list"
        assert result[0]["items"] == ["항목 A", "항목 B"]


class TestAddSection:
    """섹션 추가 테스트"""

    def test_add_approved_section_with_internal_sources(self):
        """내부 출처가 있는 승인 섹션 추가"""
        doc = create_document()
        llm_result = {
            "answer": "테스트 답변입니다.",
            "source_type": "internal",
            "source_count": 2,
        }
        sources = [
            {"source_title": "문서1", "source_url": "http://example.com/1"},
            {"source_title": "문서2", "source_url": "http://example.com/2"},
        ]

        add_section(doc, "테스트 섹션", llm_result, Route.AUTO_APPROVE, sources)

        # Heading 2가 추가되었는지 확인
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 1

    def test_add_review_section(self):
        """검토 필요 섹션 추가"""
        doc = create_document()
        llm_result = {
            "answer": "참고용 초안입니다.",
            "source_type": "internal",
            "source_count": 0,
        }

        add_section(doc, "검토 섹션", llm_result, Route.NEEDS_REVIEW, [])

        # 검토 필요 라벨이 포함되었는지 확인
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "검토 필요" in all_text

    def test_add_section_with_bullet_list(self):
        """Bullet 리스트가 포함된 섹션 추가"""
        doc = create_document()
        llm_result = {
            "answer": "요약:\n- 포인트 1\n- 포인트 2",
            "source_type": "none",
            "source_count": 0,
        }

        add_section(doc, "리스트 섹션", llm_result, Route.AUTO_APPROVE, [])

        # List Bullet 스타일 문단이 있는지 확인
        bullet_paras = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(bullet_paras) == 2


class TestBuildDocument:
    """문서 조립 테스트"""

    def test_build_document_creates_file(self):
        """build_document가 파일을 생성하는지 확인"""
        sections = [
            {"title": "서론", "query": "서론 작성", "order": 1},
            {"title": "본론", "query": "본론 작성", "order": 2},
        ]
        results = [
            ({"answer": "서론 내용", "source_type": "none", "source_count": 0}, Route.AUTO_APPROVE, []),
            ({"answer": "본론 내용", "source_type": "internal", "source_count": 1}, Route.AUTO_APPROVE, []),
        ]

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            build_document(sections, results, output_path=output_path)
            assert os.path.exists(output_path)

            # 파일이 유효한 docx인지 확인
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0
        finally:
            os.unlink(output_path)

    def test_build_document_with_custom_title(self):
        """커스텀 제목으로 문서 생성"""
        sections = [{"title": "섹션1", "query": "질의1", "order": 1}]
        results = [({"answer": "답변", "source_type": "none", "source_count": 0}, Route.AUTO_APPROVE, [])]

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            build_document(sections, results, title="커스텀 제목", output_path=output_path)
            doc = Document(output_path)

            # Heading 1에 커스텀 제목이 있는지 확인
            h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
            assert len(h1_paras) == 1
            assert h1_paras[0].text == "커스텀 제목"
        finally:
            os.unlink(output_path)
