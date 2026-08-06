"""
main.py - 파이프라인 메인 모듈

전체 grounded-report-workflow 파이프라인을 조립하고 실행합니다.

실행 흐름:
1. classify_source_type으로 섹션 분류
2. 분류 결과에 따라 검색 전략 결정 (none/internal/web/provided_data)
3. generate_section_draft로 LLM 초안 생성
4. router.route로 승인 결정
5. docx_builder로 문서 삽입

사용법:
    python -m pipeline.main

provided_data 모드:
    from pipeline.data_ingest import ingest_tabular
    chunks = ingest_tabular("data.csv")
    run_pipeline("보고서", template_hint="gambarlabs_report", provided_chunks=chunks)
"""

from typing import Callable

from pipeline.classifier import classify_source_type
from pipeline.rag_tool import search_internal_knowledge
from pipeline.web_search import search_web, _extract_topic_keywords
from pipeline.llm_writer import generate_section_draft
from pipeline.router import route, Route
from pipeline.docx_builder import (
    create_document,
    build_cover_page,
    build_toc_page,
    start_body_content,
    add_section,
    add_references_section,
    collect_references,
    save_document,
    build_from_template,
    get_template_path,
    TOC_THRESHOLD,
)
from pipeline.section_planner import plan_sections
from pipeline.templates import is_fixed_template
from pipeline.data_ingest import ProvidedChunk, filter_chunks_by_keywords

from docx import Document
from pathlib import Path


# 섹션 제목 → 템플릿 플레이스홀더 매핑
_SECTION_TO_PLACEHOLDER = {
    "개요": "OVERVIEW",
    "주요내용": "MAIN_CONTENT",
    "결론및제언": "CONCLUSION",
}

# {{FIELD}} 추론용 키워드 → 카테고리 매핑
# 주의: 키워드 중복 금지 (한 키워드는 하나의 카테고리에만 속해야 함)
_FIELD_KEYWORDS = {
    "개발": ["개발", "구현", "코딩", "프로그래밍", "API", "기능", "모듈", "코드"],
    "테스트": ["테스트", "QA", "검증", "버그", "디버깅", "품질"],
    "회의": ["회의", "미팅", "논의", "협의", "조율", "커뮤니케이션"],
    "기획": ["기획", "설계", "아키텍처", "요구사항", "정의"],  # "분석" 제거 (연구에 귀속)
    "문서화": ["문서", "매뉴얼", "가이드", "작성", "정리"],
    "배포": ["배포", "릴리즈", "운영", "인프라", "서버", "클라우드"],
    "연구": ["연구", "조사", "분석", "리서치", "POC", "실험"],  # "분석"은 여기에만
    "교육": ["교육", "학습", "세미나", "워크샵", "온보딩"],
}


def _infer_field_category(document_title: str, section_contents: list[str]) -> str:
    """
    문서 제목과 섹션 내용을 분석하여 FIELD 카테고리를 추론합니다.

    Args:
        document_title: 문서 제목
        section_contents: 섹션 내용 리스트

    Returns:
        추론된 카테고리 (예: "개발", "활동", "프로젝트 진행")
    """
    # 분석할 전체 텍스트 결합
    all_text = document_title + " " + " ".join(section_contents)
    all_text_lower = all_text.lower()

    # 각 카테고리별 매칭 점수 계산
    category_scores: dict[str, int] = {}
    for category, keywords in _FIELD_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in all_text_lower)
        if score > 0:
            category_scores[category] = score

    # 최고 점수 카테고리 반환
    if category_scores:
        best_category = max(category_scores, key=category_scores.get)
        return best_category

    # 매칭 없으면 제목에서 유형 추론
    if "보고서" in document_title:
        return "활동"
    if "제안서" in document_title:
        return "기획"

    return "프로젝트 진행"


def _generate_doc_number() -> str:
    """
    문서 번호를 생성합니다.
    형식: GBL-YYYYMMDD-순번

    Returns:
        생성된 문서 번호
    """
    from datetime import date

    today = date.today()
    date_str = today.strftime("%Y%m%d")

    # 동일 날짜 문서 수 계산 (출력 디렉토리 기준)
    # 실제로는 DB나 파일 기반으로 관리하는 것이 좋지만,
    # 여기서는 간단히 01로 고정
    seq = "01"

    return f"GBL-{date_str}-{seq}"


def _extract_date_range_from_chunks(chunks: list) -> str:
    """
    provided_chunks에서 날짜 범위를 추출합니다.

    Args:
        chunks: ProvidedChunk 리스트

    Returns:
        "YYYY년 M월 D일 ~ YYYY년 M월 D일" 형식 문자열, 없으면 빈 문자열
    """
    import re
    from datetime import datetime

    date_pattern = re.compile(r"(\d{4})-(\d{1,2})-(\d{1,2})")
    dates = []

    for chunk in chunks:
        content = chunk.get("content", "") if isinstance(chunk, dict) else ""
        matches = date_pattern.findall(content)
        for y, m, d in matches:
            try:
                dates.append(datetime(int(y), int(m), int(d)))
            except ValueError:
                pass

    if not dates:
        return ""

    min_date = min(dates)
    max_date = max(dates)

    if min_date == max_date:
        return f"{min_date.year}년 {min_date.month}월 {min_date.day}일"
    else:
        return f"{min_date.year}년 {min_date.month}월 {min_date.day}일 ~ {max_date.year}년 {max_date.month}월 {max_date.day}일"


def _build_document_from_template(
    template_path: Path,
    document_title: str,
    section_results: list[tuple],
    unique_sources: list[dict],
    slack_user_name: str | None = None,
    provided_chunks: list | None = None,
) -> Document:
    """
    템플릿 파일을 사용하여 문서를 생성합니다.

    Args:
        template_path: 템플릿 파일 경로
        document_title: 문서 제목
        section_results: 섹션 처리 결과 리스트
        unique_sources: 중복 제거된 출처 리스트
        slack_user_name: Slack 사용자 이름 (작성자)
        provided_chunks: 제공 데이터 (날짜 범위 추출용)

    Returns:
        생성된 Document 객체
    """
    from datetime import date

    today = date.today()

    # 섹션 내용 수집 (FIELD 추론용)
    section_contents = []
    for section, llm_result, decision, sources, _ in section_results:
        answer = llm_result.get("answer", "")
        if answer:
            section_contents.append(answer)

    # 날짜 범위 추출
    date_range = _extract_date_range_from_chunks(provided_chunks) if provided_chunks else ""

    # 플레이스홀더 딕셔너리 준비
    placeholders = {
        # 기본 정보
        "{{TITLE}}": document_title,
        "{{DATE}}": f"{today.year}년 {today.month}월 {today.day}일",
        "{{AUTHOR}}": slack_user_name if slack_user_name else "",
        "{{FIELD}}": _infer_field_category(document_title, section_contents),
        "{{DOC_NO}}": _generate_doc_number(),
        # 기간/프로젝트
        "{{PERIOD}}": date_range,  # CSV에서 추출한 날짜 범위
        "{{PROJECT}}": "",
        # 검토자 정보 (비워둠 - 검토 시 기입)
        "{{REVIEWER}}": "",
        "{{REVIEW_DATE}}": "",
        # 기타 플레이스홀더
        "{{PLACEHOLDER}}": "",
    }

    # 섹션 내용 매핑
    for section, llm_result, decision, sources, _ in section_results:
        title = section["title"]
        answer = llm_result.get("answer", "")
        print(f"[DEBUG] 섹션 '{title}': answer 길이={len(answer)}, decision={decision}", flush=True)

        # 섹션 제목을 플레이스홀더 키로 변환
        placeholder_key = _SECTION_TO_PLACEHOLDER.get(title)
        if placeholder_key:
            placeholders[f"{{{{{placeholder_key}}}}}"] = answer
            print(f"[DEBUG] 플레이스홀더 {{{{{placeholder_key}}}}} 설정됨", flush=True)
        else:
            print(f"[DEBUG] 플레이스홀더 매핑 없음: {title}", flush=True)

    # 참고 자료 목록 생성
    if unique_sources:
        ref_lines = []
        for i, source in enumerate(unique_sources, 1):
            src_title = source.get("source_title", "출처")
            url = source.get("source_url", "")
            if url:
                ref_lines.append(f"{i}. {src_title}\n   {url}")
            else:
                ref_lines.append(f"{i}. {src_title}")
        placeholders["{{REFERENCES}}"] = "\n".join(ref_lines)
    else:
        placeholders["{{REFERENCES}}"] = "(참고 자료 없음)"

    # 템플릿에서 문서 생성
    doc = build_from_template(template_path, placeholders)

    # 로고 삽입 (XML 레벨에서 첫 번째 테이블 앞에)
    logo_path = Path(__file__).parent.parent / "assets" / "gambalabs-logo.png"
    if logo_path.exists() and doc.tables:
        from docx.shared import Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from lxml import etree

        # 새 단락 XML 요소 생성
        new_p = OxmlElement('w:p')
        # 단락 속성 (왼쪽 정렬)
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        new_p.append(pPr)

        # 첫 번째 테이블 앞에 단락 삽입
        first_table = doc.tables[0]
        first_table._element.addprevious(new_p)

        # 삽입된 단락을 python-docx Paragraph로 래핑하여 이미지 추가
        from docx.text.paragraph import Paragraph
        logo_para = Paragraph(new_p, doc)
        run = logo_para.add_run()
        run.add_picture(str(logo_path), width=Inches(1.2))

    # 문서 전체 run 색상을 검정으로 통일 (템플릿 고정 텍스트 포함)
    from docx.shared import RGBColor
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.italic = False

    return doc


def process_section(
    section: dict,
    llm_client: Callable[[list[dict]], str] | None = None,
    classifier_client: Callable[[list[dict]], str] | None = None,
    document_title: str | None = None,
) -> tuple[dict, Route, list[dict], str]:
    """
    단일 섹션을 처리합니다.

    Args:
        section: 섹션 정보 (title, query)
        llm_client: LLM 클라이언트 (None이면 실제 Groq API 호출)
        classifier_client: 분류용 LLM 클라이언트 (None이면 실제 Groq API 호출)
        document_title: 문서 제목 (검색 쿼리에 topic 키워드 추가용)

    Returns:
        (llm_result, decision, sources, classified_type)
    """
    title = section["title"]
    query = section["query"]

    # 검색 쿼리에 topic 키워드 포함 (검색 범위 확장 방지)
    search_query = query
    if document_title and document_title not in query:
        search_query = f"{document_title} - {query}"

    # 1. classify_source_type으로 1차 분류
    classified_type = classify_source_type(
        section_title=title,
        section_query=query,
        llm_client=classifier_client,
    )

    # 2. 분류 결과에 따른 처리
    if classified_type == "none":
        # 검색 없이 LLM 호출 (llm_writer가 none 타입 전용 프롬프트 사용)
        llm_result = generate_section_draft(
            section_query=query,
            sources=[],
            source_type="none",
            llm_client=llm_client,
            topic=document_title,
        )
        decision = route(
            source_type=llm_result["source_type"],
            source_count=llm_result["source_count"],
            source_relevance=llm_result["source_relevance"],
            has_fabrication_risk=llm_result["has_fabrication_risk"],
        )
        return llm_result, decision, [], classified_type

    elif classified_type == "web":
        # 바로 웹 검색 수행 (topic 포함 쿼리)
        web_results = search_web(search_query, topic=document_title)["results"]
        sources = web_results
        llm_result = generate_section_draft(
            section_query=query,
            sources=sources,
            source_type="web",
            llm_client=llm_client,
            topic=document_title,
        )
        decision = route(
            source_type=llm_result["source_type"],
            source_count=llm_result["source_count"],
            source_relevance=llm_result["source_relevance"],
            has_fabrication_risk=llm_result["has_fabrication_risk"],
        )
        return llm_result, decision, sources, classified_type

    else:
        # classified_type == "internal"
        # 내부 RAG 검색 수행 (topic 포함 쿼리)
        internal_results = search_internal_knowledge(search_query)["results"]
        sources = internal_results
        llm_result = generate_section_draft(
            section_query=query,
            sources=sources,
            source_type="internal",
            llm_client=llm_client,
            topic=document_title,
        )

        # 결과가 없거나 source_relevance가 "low"면 웹 폴백
        if len(sources) == 0 or llm_result.get("source_relevance") == "low":
            web_results = search_web(search_query, topic=document_title)["results"]
            if web_results:
                sources = web_results
                llm_result = generate_section_draft(
                    section_query=query,
                    sources=sources,
                    source_type="web",
                    llm_client=llm_client,
                    topic=document_title,
                )

        decision = route(
            source_type=llm_result["source_type"],
            source_count=llm_result["source_count"],
            source_relevance=llm_result["source_relevance"],
            has_fabrication_risk=llm_result["has_fabrication_risk"],
        )
        return llm_result, decision, sources, classified_type


def process_section_with_provided_data(
    section: dict,
    provided_chunks: list[ProvidedChunk],
    llm_client: Callable[[list[dict]], str] | None = None,
    document_title: str | None = None,
) -> tuple[dict, Route, list[dict], str]:
    """
    제공 데이터를 사용하여 섹션을 처리합니다.

    Args:
        section: 섹션 정보 (title, query, match_keywords)
        provided_chunks: 전체 chunk 리스트 (1회성 데이터, RAG와 분리)
        llm_client: LLM 클라이언트
        document_title: 문서 제목

    Returns:
        (llm_result, decision, sources, classified_type)
    """
    title = section["title"]
    query = section["query"]

    # 1. 섹션 키워드 추출 (템플릿 정의 + 동적 추출)
    section_keywords = section.get("match_keywords", []).copy()

    # title과 query에서 추가 키워드 추출
    combined_text = f"{title} {query}"
    dynamic_keywords = _extract_topic_keywords(combined_text)
    section_keywords.extend(dynamic_keywords)

    # 중복 제거
    section_keywords = list(dict.fromkeys(section_keywords))

    # 2. 매칭되는 chunk 필터링
    matched_chunks = filter_chunks_by_keywords(
        chunks=provided_chunks,
        section_keywords=section_keywords,
        min_match=1,
    )

    print(f"  키워드: {section_keywords[:5]}{'...' if len(section_keywords) > 5 else ''} -> {len(matched_chunks)}개 매칭")

    # 3. 매칭 결과에 따른 처리
    if not matched_chunks:
        # 매칭 없음 → 개요/요약 섹션이면 전체 데이터로 요약 생성
        overview_keywords = ["개요", "요약", "소개", "배경", "목적", "overview", "summary"]
        is_overview_section = any(kw in title.lower() for kw in overview_keywords)

        if is_overview_section and provided_chunks:
            # 전체 데이터에서 상위 5개를 사용해 개요 생성
            print(f"  -> 개요 섹션: 전체 데이터({len(provided_chunks)}개)로 요약 생성")
            matched_chunks = provided_chunks[:5]
        else:
            # 다른 섹션은 근거 부족 처리
            llm_result = {
                "answer": "[근거 부족] 제공된 데이터에서 이 섹션에 해당하는 내용을 찾을 수 없습니다.",
                "source_type": "provided_data",
                "source_count": 0,
                "source_relevance": "low",
                "has_fabrication_risk": False,
            }
            decision = route(
                source_type="provided_data",
                source_count=0,
                source_relevance="low",
                has_fabrication_risk=False,
            )
            return llm_result, decision, [], "provided_data"

    # 4. LLM 호출 (매칭된 chunk를 sources로 전달)
    # chunk를 기존 sources 형식으로 변환
    sources = [
        {
            "content": chunk["content"],
            "source_title": chunk["source_title"],
            "source_url": chunk["source_url"],
            "score": chunk["score"],
        }
        for chunk in matched_chunks[:5]  # 최대 5개
    ]

    # 5. 결론/제언 섹션이면 웹 검색 추가 (인사이트, 트렌드 보강)
    conclusion_keywords = ["결론", "제언", "향후", "제안", "conclusion", "recommendation"]
    is_conclusion_section = any(kw in title.lower() for kw in conclusion_keywords)

    if is_conclusion_section and document_title:
        print(f"  -> 결론 섹션: 웹 검색으로 인사이트 보강")
        # 문서 주제 기반 웹 검색 쿼리 생성
        web_query = f"{document_title} 트렌드 인사이트 베스트프랙티스"
        web_response = search_web(web_query, max_results=3, topic=document_title)
        web_results = web_response.get("results", []) if isinstance(web_response, dict) else []

        if web_results:
            print(f"  -> 웹 검색 결과: {len(web_results)}건")
            # 웹 검색 결과를 sources에 추가
            for web_source in web_results:
                sources.append({
                    "content": web_source.get("content", ""),
                    "source_title": web_source.get("source_title", "웹 검색"),
                    "source_url": web_source.get("source_url", ""),
                    "score": web_source.get("score", 0.8),
                })

    llm_result = generate_section_draft(
        section_query=query,
        sources=sources,
        source_type="provided_data",  # 여전히 provided_data 기반 (웹은 보조)
        llm_client=llm_client,
        topic=document_title,
    )

    decision = route(
        source_type=llm_result["source_type"],
        source_count=llm_result["source_count"],
        source_relevance=llm_result["source_relevance"],
        has_fabrication_risk=llm_result["has_fabrication_risk"],
    )

    return llm_result, decision, sources, "provided_data"


def run_pipeline(
    user_request: str,
    output_path: str = "output.docx",
    template_hint: str | None = None,
    llm_client: Callable[[list[dict]], str] | None = None,
    classifier_client: Callable[[list[dict]], str] | None = None,
    planner_client: Callable[[list[dict]], str] | None = None,
    provided_chunks: list[ProvidedChunk] | None = None,
    force_fixed_template: bool = False,
    slack_user_name: str | None = None,
) -> None:
    """
    전체 파이프라인을 실행합니다.

    Args:
        user_request: 사용자 요청 (예: "클라우드 비용 최적화 제안서 작성해줘")
        output_path: 출력 파일 경로
        template_hint: 문서 유형 힌트 (예: "제안서", "기술 보고서")
        llm_client: LLM 클라이언트 (None이면 실제 Groq API 호출)
        classifier_client: 분류용 LLM 클라이언트 (None이면 실제 Groq API 호출)
        planner_client: 섹션 계획용 LLM 클라이언트 (None이면 실제 Groq API 호출)
        provided_chunks: 제공 데이터 chunk 리스트 (1회성, 메모리 한정, RAG와 분리)
        force_fixed_template: True면 고정 템플릿 강제 사용 (LLM 섹션 계획 스킵)
        slack_user_name: Slack 사용자 이름 (템플릿 {{AUTHOR}} 필드용, None이면 "[미지정]")

    Note:
        provided_chunks가 제공되면:
        - source_type="provided_data"로 분류됨
        - 기존 RAG 인덱스(rag_tool.py)와 절대 섞이지 않음
        - 실행 종료 시 메모리에서 자동 폐기됨 (GC)
    """
    print("=" * 60)
    print("Grounded Report Workflow 실행")
    print(f"요청: {user_request}")
    if template_hint:
        print(f"지정된 템플릿: {template_hint}")
    if provided_chunks:
        print(f"제공 데이터: {len(provided_chunks)}개 chunk (1회성, RAG 분리)")
    print("=" * 60)

    # 고정 템플릿 + 제공 데이터 모드 판단
    use_fixed = force_fixed_template or (
        template_hint
        and provided_chunks
        and is_fixed_template(template_hint)
    )
    print(f"[DEBUG] use_fixed={use_fixed}, force_fixed_template={force_fixed_template}, template_hint={template_hint}, provided_chunks_len={len(provided_chunks) if provided_chunks else 0}", flush=True)

    # 섹션 구성 생성 (단일 LLM 호출로 문서 제목 + 유형 감지 + 섹션 생성)
    # 고정 템플릿이면 LLM 호출 스킵
    print("\n[섹션 계획 생성 중...]")
    sections, document_title, detected_hint = plan_sections(
        user_request,
        template_hint=template_hint,
        llm_client=planner_client,
        force_fixed=use_fixed,
    )
    if document_title:
        print(f"문서 제목: {document_title}")
    if document_title:
        print(f"문서 제목: {document_title}")
    if template_hint is None and detected_hint:
        print(f"감지된 문서 유형: {detected_hint}")
    print(f"생성된 섹션: {[s['title'] for s in sections]}")

    # 문서 제목 fallback
    if not document_title:
        document_title = user_request
        for suffix in ["를 작성해줘", "을 작성해줘", "작성해줘", "해줘", "해주세요"]:
            if document_title.endswith(suffix):
                document_title = document_title[:-len(suffix)].strip()
                break

    # 템플릿 파일 경로 확인 (고정 템플릿이면서 docx_template_path가 있는 경우)
    template_file_path = get_template_path(template_hint) if template_hint else None
    use_template_file = template_file_path is not None and use_fixed
    print(f"[DEBUG] template_file_path={template_file_path}, use_template_file={use_template_file}", flush=True)

    if use_template_file:
        print(f"[템플릿 문서] {template_file_path} 사용")

    # 섹션별 처리 (출처 수집)
    all_sources = []
    section_results = []  # (section, llm_result, decision, sources, classified_type)

    for section in sections:
        title = section["title"]
        print(f"\n[처리 중] {title}...")

        # 분기: provided_data 모드 vs 기존 모드
        if provided_chunks:
            llm_result, decision, sources, classified_type = process_section_with_provided_data(
                section, provided_chunks, llm_client, document_title
            )
        else:
            llm_result, decision, sources, classified_type = process_section(
                section, llm_client, classifier_client, document_title
            )

        # 출처 수집 (참고 자료 섹션용)
        if sources:
            all_sources.append(sources)

        # 결과 저장
        section_results.append((section, llm_result, decision, sources, classified_type))

        # 진행 상황 출력
        final_source_type = llm_result.get("source_type", "unknown")
        if classified_type != final_source_type:
            print(f"[{title}] 분류: {classified_type} -> 최종: {final_source_type} -> {decision.label}")
        else:
            print(f"[{title}] -> {decision.label} ({final_source_type})")

    # 참고 자료 수집
    unique_sources = collect_references(all_sources)

    # 문서 생성 분기: 템플릿 파일 vs 일반 생성
    if use_template_file:
        # 템플릿 기반 문서 생성
        doc = _build_document_from_template(
            template_file_path,
            document_title,
            section_results,
            unique_sources,
            slack_user_name=slack_user_name,
            provided_chunks=provided_chunks,
        )
    else:
        # 일반 문서 생성
        doc = create_document()  # 빈 문서 (스타일만 적용)

        # 페이지 1: 표지
        build_cover_page(doc, title=document_title)

        # 페이지 2: 목차 (섹션 수가 충분하면)
        include_toc = len(sections) >= TOC_THRESHOLD
        if include_toc:
            build_toc_page(doc, sections=sections)

        # 페이지 3+: 본문 시작
        start_body_content(doc)

        # 섹션별 문서 추가
        for section, llm_result, decision, sources, _ in section_results:
            add_section(doc, section["title"], llm_result, decision, sources, include_inline_sources=False)

        # 참고 자료 섹션 추가
        if unique_sources:
            add_references_section(doc, unique_sources)

    if unique_sources:
        print(f"\n[참고 자료] {len(unique_sources)}건 추가됨")

    # 문서 저장
    save_document(doc, output_path)
    print("\n" + "=" * 60)
    print(f"문서 생성 완료: {output_path}")
    if provided_chunks:
        print("[제공 데이터] 1회성 데이터 메모리에서 폐기됨 (RAG 오염 없음)")
    print("=" * 60)


if __name__ == "__main__":
    import sys

    # 명령줄 인자가 있으면 사용, 없으면 기본값
    if len(sys.argv) > 1:
        request = " ".join(sys.argv[1:])
    else:
        request = "RAG 기반 문서 자동화 기술 역량 보고서를 작성해줘"

    run_pipeline(user_request=request)
