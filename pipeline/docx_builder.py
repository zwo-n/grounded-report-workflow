"""
docx_builder.py - Word 문서 생성 모듈

라우팅 결과와 LLM 응답을 기반으로 .docx 파일을 생성합니다.
콘텐츠 판단은 하지 않으며, 서식만 책임집니다.

주요 기능:
- 한글 폰트 정확한 지정 (w:eastAsia XML 레벨 설정)
- Heading 스타일 기반 목차 자동 삽입
- 실제 bullet/numbering 리스트 처리
- 표 헤더 행 음영 처리 (필요시만)
- 출처 정보 및 참고문헌 추가
- Gamba Labs 브랜드 스타일 적용
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table

from pipeline.router import Route


# =============================================================================
# 마크다운 인라인 파싱
# =============================================================================
def parse_inline_markdown(text: str) -> list[tuple[str, bool, bool]]:
    """
    마크다운 인라인 서식(볼드, 이탤릭)을 파싱합니다.

    Args:
        text: 원본 텍스트

    Returns:
        [(텍스트, bold여부, italic여부), ...] 리스트
    """
    if not text:
        return []

    result = []
    # **bold**, *italic*, ***bold+italic*** 패턴 처리
    # 패턴 우선순위: *** > ** > *
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')

    last_end = 0
    for match in pattern.finditer(text):
        # 매치 전 일반 텍스트
        if match.start() > last_end:
            result.append((text[last_end:match.start()], False, False))

        full_match = match.group(0)
        if full_match.startswith('***'):
            # bold + italic
            content = match.group(2)
            result.append((content, True, True))
        elif full_match.startswith('**'):
            # bold only
            content = match.group(3)
            result.append((content, True, False))
        else:
            # italic only
            content = match.group(4)
            result.append((content, False, True))

        last_end = match.end()

    # 남은 텍스트
    if last_end < len(text):
        result.append((text[last_end:], False, False))

    return result


def _add_formatted_text(paragraph, text: str) -> None:
    """
    마크다운 인라인 서식이 적용된 텍스트를 문단에 추가합니다.

    Args:
        paragraph: python-docx Paragraph 객체
        text: 마크다운 서식이 포함된 텍스트
    """
    segments = parse_inline_markdown(text)

    for content, is_bold, is_italic in segments:
        run = paragraph.add_run(content)
        # 폰트 먼저 설정
        _set_korean_font(run)
        run.font.size = DEFAULT_FONT_SIZE
        run.font.color.rgb = COLOR_TEXT_GRAY
        # bold/italic은 반드시 마지막에 설정 (덮어쓰기 방지)
        if is_bold:
            run.font.bold = True
        if is_italic:
            run.font.italic = True


# =============================================================================
# 하이퍼링크 헬퍼
# =============================================================================
def add_hyperlink(paragraph, url: str, text: str, color: RGBColor | None = None) -> None:
    """
    문단에 클릭 가능한 하이퍼링크를 추가합니다.

    Args:
        paragraph: python-docx Paragraph 객체
        url: 링크 URL
        text: 표시할 텍스트
        color: 링크 색상 (None이면 기본 파란색)
    """
    # 파트에서 relationship 생성
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # w:hyperlink 요소 생성
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # w:r (run) 요소 생성
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 링크 스타일 (파란색 + 밑줄)
    c = OxmlElement('w:color')
    if color:
        c.set(qn('w:val'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
    else:
        c.set(qn('w:val'), '0563C1')  # 기본 링크 파란색
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # 폰트 설정
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), DEFAULT_FONT_NAME)
    rFonts.set(qn('w:ascii'), DEFAULT_FONT_NAME)
    rFonts.set(qn('w:hAnsi'), DEFAULT_FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)

    new_run.append(rPr)

    # 텍스트 추가
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# =============================================================================
# 번호 매기기 리스트 (독립 numId)
# =============================================================================
def _get_fresh_numbering_id(doc: Document) -> int:
    """
    문서에서 새로운 독립적인 numbering ID를 생성합니다.
    각 번호 리스트가 1번부터 시작하도록 합니다.

    Args:
        doc: python-docx Document 객체

    Returns:
        새로운 numId
    """
    # numbering.xml 파트 가져오기
    numbering_part = doc.part.numbering_part

    if numbering_part is None:
        # numbering 파트가 없으면 생성 (첫 번째 리스트 추가 시 자동 생성됨)
        temp_para = doc.add_paragraph("temp", style='List Number')
        doc._body._element.remove(temp_para._element)
        numbering_part = doc.part.numbering_part

    numbering_xml = numbering_part._element

    # 기존 abstractNum 찾기 (List Number 스타일용)
    # namespace 없이 직접 접근
    abstract_num_id = "0"
    for child in numbering_xml:
        if child.tag.endswith('abstractNum'):
            # abstractNumId 속성 가져오기 (namespace 포함된 키로)
            for key in child.attrib:
                if key.endswith('abstractNumId'):
                    abstract_num_id = child.attrib[key]
                    break
            if abstract_num_id != "0":
                break

    # 기존 num 요소들의 최대 numId 찾기
    max_num_id = 0
    for child in numbering_xml:
        if child.tag.endswith('num'):
            for key in child.attrib:
                if key.endswith('numId'):
                    try:
                        num_id_val = int(child.attrib[key])
                        max_num_id = max(max_num_id, num_id_val)
                    except ValueError:
                        pass
                    break

    # 새로운 numId 생성 (기존 최대값 + 1)
    new_num_id = max_num_id + 1

    # 새로운 num 요소 생성
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(new_num_id))

    abstract_num_id_elem = OxmlElement('w:abstractNumId')
    abstract_num_id_elem.set(qn('w:val'), abstract_num_id)
    num.append(abstract_num_id_elem)

    # lvlOverride로 시작 번호 1로 강제 설정
    lvl_override = OxmlElement('w:lvlOverride')
    lvl_override.set(qn('w:ilvl'), '0')

    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), '1')
    lvl_override.append(start_override)

    num.append(lvl_override)

    # numbering.xml에 추가
    numbering_xml.append(num)

    return new_num_id


def _add_numbered_list_fresh(doc: Document, items: list[str]) -> None:
    """
    1번부터 시작하는 독립적인 번호 리스트를 추가합니다.

    Args:
        doc: python-docx Document 객체
        items: 리스트 아이템 문자열 리스트
    """
    if not items:
        return

    # 새로운 numId 발급
    num_id = _get_fresh_numbering_id(doc)

    for item in items:
        para = doc.add_paragraph()
        para.style = doc.styles['List Number']
        para.paragraph_format.space_after = Pt(4)

        # numId 재설정
        pPr = para._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            numPr = OxmlElement('w:numPr')
            pPr.append(numPr)

        # 기존 numId 제거 후 새 numId 설정
        existing_numId = numPr.find(qn('w:numId'))
        if existing_numId is not None:
            numPr.remove(existing_numId)

        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), str(num_id))
        numPr.append(numId_elem)

        # ilvl 설정 (레벨 0)
        existing_ilvl = numPr.find(qn('w:ilvl'))
        if existing_ilvl is not None:
            numPr.remove(existing_ilvl)

        ilvl_elem = OxmlElement('w:ilvl')
        ilvl_elem.set(qn('w:val'), '0')
        numPr.insert(0, ilvl_elem)

        # 마크다운 서식 적용하여 텍스트 추가
        _add_formatted_text(para, item)


# =============================================================================
# 기본 로고 경로
# =============================================================================
DEFAULT_LOGO_PATH = Path(__file__).parent.parent / "assets" / "gambalabs-logo.png"


# =============================================================================
# Gamba Labs 브랜드 컬러 시스템
# =============================================================================
# 주요 색상 (Primary)
COLOR_PRIMARY = RGBColor(0x00, 0x38, 0x99)       # #003899 - 진한 파랑
COLOR_PRIMARY_DARK = RGBColor(0x0c, 0x13, 0x31)  # #0c1331 - 거의 검은 남색

# 보조 색상 (Secondary)
COLOR_ACCENT = RGBColor(0x20, 0xae, 0xe5)        # #20aee5 - 밝은 시안

# 중성 색상 (Neutral)
COLOR_LIGHT_BG = RGBColor(0xf7, 0xf7, 0xfa)      # #f7f7fa - 밝은 회색 배경
COLOR_TEXT_GRAY = RGBColor(0x3c, 0x3c, 0x3c)     # #3c3c3c - 진한 회색 텍스트
COLOR_TEXT_MUTED = RGBColor(0x6c, 0x6c, 0x6c)    # #6c6c6c - 연한 회색 텍스트

# 경고/알림 색상
COLOR_WARNING_BG = RGBColor(0xff, 0xf3, 0xcd)    # #fff3cd - 경고 배경 (연한 노랑)
COLOR_WARNING_TEXT = RGBColor(0x85, 0x6d, 0x04)  # #856d04 - 경고 텍스트 (진한 노랑)
COLOR_WARNING_BORDER = "E0A800"                   # 경고 테두리 (HEX)

# =============================================================================
# 폰트 및 레이아웃 설정
# =============================================================================
DEFAULT_FONT_NAME = "맑은 고딕"
DEFAULT_FONT_SIZE = Pt(11)

# 헤딩 스타일 설정 (크기, 굵기, 색상)
HEADING_STYLES = {
    1: {"size": Pt(24), "bold": True, "color": COLOR_PRIMARY_DARK, "space_before": Pt(0), "space_after": Pt(12)},
    2: {"size": Pt(16), "bold": True, "color": COLOR_PRIMARY, "space_before": Pt(18), "space_after": Pt(8)},
    3: {"size": Pt(13), "bold": True, "color": COLOR_PRIMARY, "space_before": Pt(12), "space_after": Pt(6)},
}

# 본문 줄간격 설정
LINE_SPACING = 1.5  # 1.5줄
PARAGRAPH_SPACE_AFTER = Pt(8)

# 목차 삽입 기준 섹션 수
TOC_THRESHOLD = 5


# =============================================================================
# 폰트 설정 헬퍼 함수
# =============================================================================
def _set_korean_font(run, font_name: str = DEFAULT_FONT_NAME) -> None:
    """
    Run에 한글 폰트를 정확하게 설정합니다.

    w:eastAsia 속성까지 XML 레벨에서 설정하여
    run.font.name만으로는 한글 폰트가 적용되지 않는 문제를 해결합니다.

    Args:
        run: python-docx Run 객체
        font_name: 적용할 폰트명
    """
    run.font.name = font_name
    # w:rFonts 요소의 w:eastAsia 속성 설정
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def _set_style_korean_font(style, font_name: str = DEFAULT_FONT_NAME) -> None:
    """
    스타일에 한글 폰트를 설정합니다.

    Args:
        style: python-docx Style 객체
        font_name: 폰트명
    """
    style.font.name = font_name
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


# =============================================================================
# 헤더/로고 관련 함수
# =============================================================================
def _add_header_with_logo(doc: Document, logo_path: str | Path | None = None) -> bool:
    """
    문서 헤더에 로고 이미지를 좌상단에 삽입합니다.

    Args:
        doc: python-docx Document 객체
        logo_path: 로고 이미지 경로 (None이면 기본 경로 사용)

    Returns:
        로고 삽입 성공 여부
    """
    if logo_path is None:
        logo_path = DEFAULT_LOGO_PATH

    logo_path = Path(logo_path)

    if not logo_path.exists():
        print(f"[경고] 로고 파일을 찾을 수 없습니다: {logo_path}")
        return False

    # 첫 번째 섹션의 헤더 가져오기
    section = doc.sections[0]
    header = section.header

    # 헤더 문단에 로고 삽입 (좌상단 정렬)
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = header_para.add_run()
    # 로고 크기: 너비 1.5인치 (높이는 비율 유지)
    run.add_picture(str(logo_path), width=Inches(1.5))

    return True


def _add_metadata_block(
    doc: Document,
    author: str | None = None,
    date: str | None = None,
) -> None:
    """
    제목 아래 메타데이터 블록(작성자, 작성일)을 추가합니다.

    Args:
        doc: python-docx Document 객체
        author: 작성자 (None이면 플레이스홀더)
        date: 작성일 (None이면 플레이스홀더)
    """
    # 값이 없으면 명시적 플레이스홀더 사용 (임의 추정 금지)
    author_text = author if author else "[작성자]"
    date_text = date if date else "[작성일]"

    metadata_text = f"작성자: {author_text}    |    작성일: {date_text}"

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)

    run = para.add_run(metadata_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)  # #666666
    _set_korean_font(run)


def _add_separator_line(doc: Document) -> None:
    """
    구분선(하단 테두리가 있는 빈 문단)을 추가합니다.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)

    # 하단 테두리만 추가 (연회색 실선)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')  # 연회색
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_document_header(
    doc: Document,
    title: str,
    logo_path: str | Path | None = None,
    author: str | None = None,
    date: str | None = None,
) -> None:
    """
    문서 공식 레이아웃(로고, 제목, 메타데이터)을 추가합니다.

    Args:
        doc: python-docx Document 객체
        title: 문서 제목
        logo_path: 로고 이미지 경로 (None이면 기본 로고, False면 로고 없음)
        author: 작성자 (None이면 플레이스홀더)
        date: 작성일 (None이면 플레이스홀더)
    """
    # 1. 로고 삽입 (헤더 영역)
    if logo_path is not False:
        _add_header_with_logo(doc, logo_path)

    # 2. 제목 (Heading 1)
    heading = doc.add_heading(title, level=1)
    style_config = HEADING_STYLES[1]
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = style_config["space_after"]

    for run in heading.runs:
        run.font.size = style_config["size"]
        run.font.bold = style_config["bold"]
        run.font.color.rgb = style_config["color"]
        _set_korean_font(run)

    # 3. 메타데이터 블록 (작성자, 작성일)
    _add_metadata_block(doc, author, date)


# =============================================================================
# 페이지 구조 함수 (표지, 목차, 참고자료)
# =============================================================================
def build_cover_page(
    doc: Document,
    title: str,
    logo_path: str | Path | None = None,
    author: str | None = None,
    date: str | None = None,
) -> None:
    """
    표지 페이지를 생성합니다 (1페이지).
    로고, 제목, 작성자/작성일을 포함합니다.

    Args:
        doc: python-docx Document 객체
        title: 문서 제목
        logo_path: 로고 이미지 경로 (None이면 기본 로고)
        author: 작성자 (None이면 플레이스홀더)
        date: 작성일 (None이면 플레이스홀더)
    """
    # 1. 로고 삽입 (헤더 영역, 좌상단)
    _add_header_with_logo(doc, logo_path)

    # 2. 여백 추가 (상단 여백)
    for _ in range(3):
        doc.add_paragraph()

    # 3. 제목 (대형, 중앙 정렬)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(24)

    title_run = title_para.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY_DARK
    _set_korean_font(title_run)

    # 4. 여백 추가
    for _ in range(4):
        doc.add_paragraph()

    # 5. 메타데이터 (중앙 정렬)
    author_text = author if author else "[작성자]"
    date_text = date if date else "[작성일]"
    metadata_text = f"작성자: {author_text}    |    작성일: {date_text}"

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_before = Pt(0)
    meta_para.paragraph_format.space_after = Pt(8)

    meta_run = meta_para.add_run(metadata_text)
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = COLOR_TEXT_MUTED
    _set_korean_font(meta_run)


def build_toc_page(doc: Document, sections: list[dict] | None = None) -> None:
    """
    목차 페이지를 생성합니다 (2페이지).
    Word TOC 필드를 삽입하여 Word에서 업데이트하면 자동으로 페이지 번호가 채워집니다.

    Args:
        doc: python-docx Document 객체
        sections: 섹션 목록 (정적 목차 미리보기용, None이면 동적 TOC만)
    """
    # 페이지 브레이크
    doc.add_page_break()

    # "목차" 제목
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_before = Pt(24)
    toc_title.paragraph_format.space_after = Pt(18)

    toc_run = toc_title.add_run("목차")
    toc_run.font.size = Pt(18)
    toc_run.font.bold = True
    toc_run.font.color.rgb = COLOR_PRIMARY
    _set_korean_font(toc_run)

    # Word TOC 필드 삽입 (업데이트하면 점선+페이지번호 자동 생성)
    _add_toc_with_preview(doc, sections)

    # 목차 후 여백
    doc.add_paragraph()


def _add_toc_with_preview(doc: Document, sections: list[dict] | None = None) -> None:
    """
    TOC 필드와 미리보기 텍스트를 함께 삽입합니다.
    Word에서 필드 업데이트 전까지 미리보기 텍스트가 표시됩니다.

    Args:
        doc: python-docx Document 객체
        sections: 미리보기용 섹션 목록
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # TOC 필드 시작
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # \o "1-3": Heading 1~3 포함, \h: 하이퍼링크, \z: 탭리더 숨김 해제, \u: 단락 서식 사용
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    r_element = run._r
    r_element.append(fldChar_begin)
    r_element.append(instrText)
    r_element.append(fldChar_separate)

    # 미리보기 텍스트 (섹션 목록이 있으면 표시)
    # 표지(1p) + 목차(2p) 이후 본문은 3페이지부터 시작
    if sections:
        start_page = 3
        for i, section in enumerate(sections, 1):
            title = section.get("title", f"섹션 {i}")
            preview_para = doc.add_paragraph()
            # 예상 페이지 번호 (실제로는 Word에서 필드 업데이트 필요)
            estimated_page = str(start_page + i - 1)
            _add_toc_entry_with_dots(preview_para, f"{i}. {title}", estimated_page)

        # 참고 자료 항목
        ref_para = doc.add_paragraph()
        ref_page = str(start_page + len(sections))
        _add_toc_entry_with_dots(ref_para, f"{len(sections) + 1}. 참고 자료", ref_page)

        # 안내 문구
        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_para.add_run("(Word에서 목차를 우클릭하여 '필드 업데이트'를 선택하면 페이지 번호가 표시됩니다)")
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = COLOR_TEXT_MUTED
        _set_korean_font(note_run)

    # TOC 필드 종료
    end_para = doc.add_paragraph()
    end_run = end_para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(fldChar_end)


def _add_toc_entry_with_dots(paragraph, title: str, page_num: str) -> None:
    """
    점 리더와 페이지 번호가 있는 목차 항목을 추가합니다.

    Args:
        paragraph: python-docx Paragraph 객체
        title: 섹션 제목
        page_num: 페이지 번호 (빈 문자열이면 생략)
    """
    # 탭 스톱 설정 (우측 정렬, 점 리더)
    pPr = paragraph._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9360')  # 6.5인치 (우측 여백)
    tabs.append(tab)
    pPr.append(tabs)

    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Inches(0.3)

    # 제목
    title_run = paragraph.add_run(title)
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = COLOR_PRIMARY_DARK
    _set_korean_font(title_run)

    # 탭 + 페이지 번호
    if page_num:
        tab_run = paragraph.add_run('\t')
        page_run = paragraph.add_run(page_num)
        page_run.font.size = Pt(11)
        page_run.font.color.rgb = COLOR_TEXT_MUTED
        _set_korean_font(page_run)


def start_body_content(doc: Document) -> None:
    """
    본문 시작 페이지로 이동합니다 (3페이지부터).
    페이지 브레이크를 추가합니다.

    Args:
        doc: python-docx Document 객체
    """
    doc.add_page_break()


def collect_references(all_sources: list[list[dict]]) -> list[dict]:
    """
    모든 섹션의 출처를 수집하고 중복을 제거합니다.

    Args:
        all_sources: 섹션별 출처 리스트의 리스트

    Returns:
        중복 제거된 출처 리스트
    """
    seen_urls = set()
    unique_sources = []

    for sources in all_sources:
        for source in sources:
            url = source.get("source_url", "")
            title = source.get("source_title", "")

            # URL 기준으로 중복 체크 (URL 없으면 제목 기준)
            key = url if url else title
            if key and key not in seen_urls:
                seen_urls.add(key)
                unique_sources.append(source)

    return unique_sources


def add_references_section(doc: Document, sources: list[dict]) -> None:
    """
    참고 자료 섹션을 문서 끝에 추가합니다.
    URL은 클릭 가능한 하이퍼링크로 표시됩니다.

    Args:
        doc: python-docx Document 객체
        sources: 출처 리스트 (collect_references 결과)
    """
    if not sources:
        return

    # "참고 자료" 섹션 제목 (Heading 2)
    heading = doc.add_heading("참고 자료", level=2)
    style_config = HEADING_STYLES[2]

    heading.paragraph_format.space_before = Pt(24)
    heading.paragraph_format.space_after = style_config["space_after"]

    for run in heading.runs:
        run.font.size = style_config["size"]
        run.font.bold = style_config["bold"]
        run.font.color.rgb = style_config["color"]
        _set_korean_font(run)

    # 출처 목록 (독립적인 번호 리스트 - 1번부터 시작)
    num_id = _get_fresh_numbering_id(doc)

    for i, source in enumerate(sources, 1):
        title = source.get("source_title", "출처")
        url = source.get("source_url", "")

        para = doc.add_paragraph()
        para.style = doc.styles['List Number']
        para.paragraph_format.space_after = Pt(6)

        # numId 재설정 (독립적인 번호 매기기)
        pPr = para._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            numPr = OxmlElement('w:numPr')
            pPr.append(numPr)

        existing_numId = numPr.find(qn('w:numId'))
        if existing_numId is not None:
            numPr.remove(existing_numId)

        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), str(num_id))
        numPr.append(numId_elem)

        existing_ilvl = numPr.find(qn('w:ilvl'))
        if existing_ilvl is not None:
            numPr.remove(existing_ilvl)

        ilvl_elem = OxmlElement('w:ilvl')
        ilvl_elem.set(qn('w:val'), '0')
        numPr.insert(0, ilvl_elem)

        # 제목 추가
        title_run = para.add_run(title)
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = COLOR_TEXT_GRAY
        _set_korean_font(title_run)

        # URL이 있으면 줄바꿈 후 하이퍼링크 추가
        if url:
            # 내부 문서(rag://...)는 하이퍼링크 없이 일반 텍스트로
            if url.startswith("rag://") or url.startswith("internal/"):
                para.add_run("\n   ")
                url_run = para.add_run(url)
                url_run.font.size = Pt(9)
                url_run.font.color.rgb = COLOR_TEXT_MUTED
                _set_korean_font(url_run)
            else:
                # 외부 URL은 클릭 가능한 하이퍼링크로
                para.add_run("\n   ")
                add_hyperlink(para, url, url)


def _set_paragraph_shading(paragraph, fill_color: str) -> None:
    """
    문단에 배경색(shading)을 설정합니다.

    Args:
        paragraph: python-docx Paragraph 객체
        fill_color: HEX 색상 코드 (# 제외)
    """
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)


def _set_paragraph_borders(paragraph, color: str, size: int = 4) -> None:
    """
    문단에 테두리를 설정합니다.

    Args:
        paragraph: python-docx Paragraph 객체
        color: HEX 색상 코드 (# 제외)
        size: 테두리 두께 (8분의 1 포인트 단위)
    """
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), color)
        pBdr.append(border)

    pPr.append(pBdr)


def _apply_body_paragraph_format(paragraph) -> None:
    """
    본문 문단에 표준 서식(줄간격, 여백)을 적용합니다.

    Args:
        paragraph: python-docx Paragraph 객체
    """
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = PARAGRAPH_SPACE_AFTER


# =============================================================================
# 목차(TOC) 관련 함수
# =============================================================================
def _create_toc_paragraph() -> OxmlElement:
    """
    목차 필드를 포함한 문단 XML 요소를 생성합니다.

    Returns:
        목차 필드가 포함된 w:p 요소
    """
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_separate)
    r.append(fldChar_end)

    p.append(r)
    return p


def _create_toc_title_paragraph(font_name: str = DEFAULT_FONT_NAME) -> OxmlElement:
    """
    "목차" 제목 문단 XML 요소를 생성합니다.

    Args:
        font_name: 폰트명

    Returns:
        "목차" 제목이 포함된 w:p 요소
    """
    p = OxmlElement('w:p')

    # 문단 속성 (가운데 정렬, 상하 여백)
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    # 상단 여백
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')  # 12pt
    spacing.set(qn('w:after'), '120')   # 6pt
    pPr.append(spacing)

    p.append(pPr)

    # Run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    b = OxmlElement('w:b')
    rPr.append(b)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '32')  # 16pt
    rPr.append(sz)

    # 색상 적용 (브랜드 컬러)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '003899')
    rPr.append(color)

    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = "목차"
    r.append(t)

    p.append(r)
    return p


def _add_toc(doc: Document) -> None:
    """
    문서에 목차(TOC)를 삽입합니다.

    Args:
        doc: python-docx Document 객체
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar_begin)
    r_element.append(instrText)
    r_element.append(fldChar_separate)
    r_element.append(fldChar_end)


def _insert_toc_after_title(doc: Document) -> None:
    """
    문서 제목(Heading 1) 바로 다음에 목차를 삽입합니다.

    Args:
        doc: python-docx Document 객체
    """
    body = doc._body._element

    insert_index = 0
    for i, child in enumerate(body):
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                    insert_index = i + 1
                    break

    toc_title = _create_toc_title_paragraph()
    body.insert(insert_index, toc_title)

    toc_field = _create_toc_paragraph()
    body.insert(insert_index + 1, toc_field)

    empty_p = OxmlElement('w:p')
    body.insert(insert_index + 2, empty_p)


def _count_heading2(doc: Document) -> int:
    """
    문서 내 Heading 2 수를 카운트합니다.

    Args:
        doc: python-docx Document 객체

    Returns:
        Heading 2 개수
    """
    count = 0
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 2':
            count += 1
    return count


# =============================================================================
# 리스트 관련 함수
# =============================================================================
def _add_bullet_list(doc: Document, items: list[str]) -> None:
    """
    실제 bullet 리스트를 문서에 추가합니다.
    **볼드**, *이탤릭* 등 인라인 마크다운 서식도 적용됩니다.

    Args:
        doc: python-docx Document 객체
        items: 리스트 아이템 문자열 리스트
    """
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_after = Pt(4)
        # 마크다운 인라인 서식 적용
        _add_formatted_text(para, item)


def _add_numbered_list(doc: Document, items: list[str]) -> None:
    """
    실제 numbered 리스트를 문서에 추가합니다.

    Args:
        doc: python-docx Document 객체
        items: 리스트 아이템 문자열 리스트
    """
    for item in items:
        para = doc.add_paragraph(item, style='List Number')
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.size = DEFAULT_FONT_SIZE
            run.font.color.rgb = COLOR_TEXT_GRAY
            _set_korean_font(run)


# =============================================================================
# 표(Table) 관련 함수
# =============================================================================
def _add_table_with_header(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
) -> Table:
    """
    헤더 행에 브랜드 스타일이 적용된 표를 추가합니다.

    Args:
        doc: python-docx Document 객체
        headers: 헤더 열 리스트
        rows: 데이터 행 리스트

    Returns:
        생성된 Table 객체
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행 설정 (브랜드 컬러 적용)
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text

        # 헤더 셀 음영 (브랜드 Primary Dark 색상)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003899')  # 브랜드 주요 색상
        cell._tc.get_or_add_tcPr().append(shading)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # 흰색 텍스트
                run.font.size = Pt(10)
                _set_korean_font(run)

    # 데이터 행 설정
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text

            # 짝수 행 배경색 (zebra stripe)
            if row_idx % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F7F7FA')
                cell._tc.get_or_add_tcPr().append(shading)

            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = COLOR_TEXT_GRAY
                    _set_korean_font(run)

    return table


# =============================================================================
# 마크다운 표 파싱
# =============================================================================
def _is_table_separator(line: str) -> bool:
    """마크다운 표 구분선인지 확인합니다."""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    # |---|---| 또는 |:---|:---:| 형태
    inner = stripped[1:-1]
    cells = inner.split('|')
    for cell in cells:
        cell = cell.strip()
        # 구분선 셀: -만 포함하거나 :- -: :-: 형태
        if not re.match(r'^:?-+:?$', cell):
            return False
    return True


def _parse_table_row(line: str) -> list[str]:
    """마크다운 표 행을 파싱하여 셀 리스트를 반환합니다."""
    stripped = line.strip()
    if stripped.startswith('|'):
        stripped = stripped[1:]
    if stripped.endswith('|'):
        stripped = stripped[:-1]
    cells = [cell.strip() for cell in stripped.split('|')]
    return cells


def _parse_markdown_table(lines: list[str], start_idx: int) -> tuple[dict | None, int]:
    """
    마크다운 표를 파싱합니다.

    Args:
        lines: 전체 라인 리스트
        start_idx: 시작 인덱스

    Returns:
        (표 구조 dict, 다음 처리할 인덱스) 또는 (None, start_idx) 파싱 실패시
    """
    if start_idx >= len(lines):
        return None, start_idx

    first_line = lines[start_idx].strip()

    # 첫 번째 줄이 표 행인지 확인 (| 로 시작)
    if not first_line.startswith('|'):
        return None, start_idx

    # 최소 2줄 필요 (헤더 + 구분선)
    if start_idx + 1 >= len(lines):
        return None, start_idx

    second_line = lines[start_idx + 1].strip()
    if not _is_table_separator(second_line):
        return None, start_idx

    # 헤더 행 파싱
    headers = _parse_table_row(first_line)
    if not headers:
        return None, start_idx

    # 데이터 행 파싱
    rows = []
    current_idx = start_idx + 2  # 구분선 다음부터

    while current_idx < len(lines):
        line = lines[current_idx].strip()
        # 빈 줄이거나 표 행이 아니면 종료
        if not line or not line.startswith('|'):
            break
        # 구분선은 무시 (다중 구분선 허용)
        if _is_table_separator(line):
            current_idx += 1
            continue
        row_cells = _parse_table_row(line)
        # 열 개수 맞추기
        while len(row_cells) < len(headers):
            row_cells.append('')
        rows.append(row_cells[:len(headers)])
        current_idx += 1

    return {
        "type": "table",
        "headers": headers,
        "rows": rows
    }, current_idx


# =============================================================================
# 콘텐츠 파싱 및 구조화
# =============================================================================
def _parse_content_structure(text: str) -> list[dict]:
    """
    텍스트를 파싱하여 구조화된 콘텐츠로 변환합니다.
    마크다운 표, 리스트, 일반 문단을 인식합니다.

    Args:
        text: 원본 텍스트

    Returns:
        구조화된 콘텐츠 리스트
    """
    if not text:
        return []

    lines = text.strip().split('\n')
    result = []
    current_list = []
    current_list_type = None
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 마크다운 표 감지
        if stripped.startswith('|'):
            # 현재 리스트 종료
            if current_list:
                result.append({
                    "type": current_list_type + "_list",
                    "items": current_list
                })
                current_list = []
                current_list_type = None

            table_data, next_idx = _parse_markdown_table(lines, i)
            if table_data:
                result.append(table_data)
                i = next_idx
                continue
            # 표 파싱 실패시 일반 문단으로 처리

        # Bullet 리스트
        if stripped.startswith('- ') or stripped.startswith('* '):
            if current_list_type != 'bullet':
                if current_list:
                    result.append({
                        "type": current_list_type + "_list",
                        "items": current_list
                    })
                current_list = []
                current_list_type = 'bullet'
            current_list.append(stripped[2:])
            i += 1
            continue

        # Numbered 리스트
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
            if current_list_type != 'numbered':
                if current_list:
                    result.append({
                        "type": current_list_type + "_list",
                        "items": current_list
                    })
                current_list = []
                current_list_type = 'numbered'
            current_list.append(stripped[2:].strip())
            i += 1
            continue

        # 리스트 종료
        if current_list:
            result.append({
                "type": current_list_type + "_list",
                "items": current_list
            })
            current_list = []
            current_list_type = None

        if not stripped:
            i += 1
            continue

        result.append({"type": "paragraph", "content": stripped})
        i += 1

    if current_list:
        result.append({
            "type": current_list_type + "_list",
            "items": current_list
        })

    return result


def _add_structured_content(doc: Document, text: str) -> None:
    """
    구조화된 콘텐츠를 문서에 추가합니다.
    마크다운 표, 리스트, 일반 문단을 렌더링합니다.
    **볼드**, *이탤릭* 등 인라인 마크다운 서식도 적용됩니다.

    Args:
        doc: python-docx Document 객체
        text: 원본 텍스트
    """
    structures = _parse_content_structure(text)

    for item in structures:
        if item["type"] == "paragraph":
            para = doc.add_paragraph()
            _apply_body_paragraph_format(para)
            # 마크다운 인라인 서식 적용
            _add_formatted_text(para, item["content"])
        elif item["type"] == "bullet_list":
            _add_bullet_list(doc, item["items"])
        elif item["type"] == "numbered_list":
            # 독립적인 번호 매기기 (1번부터 시작)
            _add_numbered_list_fresh(doc, item["items"])
        elif item["type"] == "table":
            # 마크다운 표 렌더링
            _add_table_with_header(doc, item["headers"], item["rows"])
            # 표 후 여백 추가
            doc.add_paragraph()


# =============================================================================
# 출처 포맷팅
# =============================================================================
def _format_internal_sources(sources: list[dict]) -> tuple[str, str]:
    """
    내부 RAG 출처를 포맷팅합니다.

    Args:
        sources: 검색 결과 리스트

    Returns:
        (인라인 출처 문자열, 참고문서 목록 문자열)
    """
    if not sources:
        return "", ""

    titles = [s.get("source_title", "문서") for s in sources]
    unique_titles = list(dict.fromkeys(titles))

    inline = f"[출처: {', '.join(unique_titles)}]"

    ref_lines = []
    for s in sources:
        title = s.get("source_title", "문서")
        url = s.get("source_url", "")
        if url:
            ref_lines.append(f"- {title} ({url})")
        else:
            ref_lines.append(f"- {title}")

    ref_list = "참고문서:\n" + "\n".join(ref_lines)
    return inline, ref_list


def _format_web_sources(sources: list[dict]) -> str:
    """
    웹 검색 출처를 포맷팅합니다.

    Args:
        sources: 검색 결과 리스트

    Returns:
        인라인 출처 문자열
    """
    if not sources:
        return ""

    parts = []
    for s in sources:
        title = s.get("source_title", "웹페이지")
        url = s.get("source_url", "")
        if url:
            parts.append(f"{title} ({url})")
        else:
            parts.append(title)

    return f"[출처: {', '.join(parts)}]"


# =============================================================================
# 섹션 추가 함수
# =============================================================================
def add_section(
    doc: Document,
    section_title: str,
    llm_result: dict,
    decision: Route,
    sources: list[dict],
    include_inline_sources: bool = False,
) -> None:
    """
    문서에 섹션을 추가합니다.

    Args:
        doc: python-docx Document 객체
        section_title: 섹션 제목
        llm_result: LLM 응답 결과
        decision: 라우팅 결과
        sources: 검색 결과 리스트
        include_inline_sources: 섹션 내 인라인 출처 표시 여부 (기본: False)
    """
    # 섹션 제목 (Heading 2) - 브랜드 스타일 적용
    heading = doc.add_heading(section_title, level=2)
    style_config = HEADING_STYLES[2]

    heading.paragraph_format.space_before = style_config["space_before"]
    heading.paragraph_format.space_after = style_config["space_after"]

    for run in heading.runs:
        run.font.size = style_config["size"]
        run.font.bold = style_config["bold"]
        run.font.color.rgb = style_config["color"]
        _set_korean_font(run)

    answer = llm_result.get("answer", "")
    source_type = llm_result.get("source_type", "none")

    if decision == Route.AUTO_APPROVE:
        _add_approved_section(doc, answer, source_type, sources, include_inline_sources)
    else:
        _add_review_section(doc, answer, llm_result.get("source_count", 0), decision)


def _add_approved_section(
    doc: Document,
    answer: str,
    source_type: str,
    sources: list[dict],
    include_inline_sources: bool = False,
) -> None:
    """
    자동 승인된 섹션을 추가합니다.

    Args:
        doc: Document 객체
        answer: LLM 생성 답변
        source_type: 출처 유형
        sources: 검색 결과 리스트
        include_inline_sources: 인라인 출처 표시 여부
    """
    if answer:
        _add_structured_content(doc, answer)

    # 인라인 출처 추가 (옵션)
    if not include_inline_sources:
        return

    # 출처 추가 (브랜드 Accent 색상 적용)
    if source_type == "internal":
        inline_ref, ref_list = _format_internal_sources(sources)
        if inline_ref:
            ref_para = doc.add_paragraph(inline_ref)
            ref_para.paragraph_format.space_before = Pt(8)
            for run in ref_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_ACCENT  # 브랜드 시안 색상
                _set_korean_font(run)

        if ref_list:
            ref_para = doc.add_paragraph(ref_list)
            for run in ref_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_TEXT_MUTED
                _set_korean_font(run)

    elif source_type == "web":
        inline_ref = _format_web_sources(sources)
        if inline_ref:
            ref_para = doc.add_paragraph(inline_ref)
            ref_para.paragraph_format.space_before = Pt(8)
            for run in ref_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_ACCENT
                _set_korean_font(run)


def _add_review_section(
    doc: Document,
    answer: str,
    source_count: int,
    decision: Route,
) -> None:
    """
    검토 필요 섹션을 추가합니다.
    배지/박스 스타일로 눈에 띄게 표시합니다.

    Args:
        doc: Document 객체
        answer: LLM 생성 답변 (참고용)
        source_count: 출처 개수
        decision: 라우팅 결과
    """
    # 경고 배지 스타일 (배경색 + 테두리)
    if source_count == 0:
        reason = "근거 문서 없음"
    else:
        reason = "근거 부족 또는 신뢰도 낮음"

    notice_text = f"  {decision.label}  |  {reason}"
    notice = doc.add_paragraph(notice_text)

    # 배지 스타일 적용
    notice.paragraph_format.space_before = Pt(8)
    notice.paragraph_format.space_after = Pt(8)
    notice.paragraph_format.left_indent = Inches(0.2)
    notice.paragraph_format.right_indent = Inches(0.2)

    # 배경색 (연한 노랑)
    _set_paragraph_shading(notice, 'FFF3CD')
    # 테두리 (진한 노랑)
    _set_paragraph_borders(notice, 'E0A800', size=6)

    for run in notice.runs:
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WARNING_TEXT
        _set_korean_font(run)

    # 참고용 초안
    if answer:
        draft_label = doc.add_paragraph("참고용 초안")
        draft_label.paragraph_format.space_before = Pt(12)
        draft_label.paragraph_format.space_after = Pt(4)
        for run in draft_label.runs:
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = COLOR_TEXT_MUTED
            _set_korean_font(run)

        _add_structured_content(doc, answer)

        # 초안 영역 시각적 구분 (연한 배경)
        # 마지막 추가된 문단들에 이탤릭 적용
        for para in doc.paragraphs[-1:]:
            for run in para.runs:
                run.italic = True


# =============================================================================
# 문서 생성 및 저장
# =============================================================================
def create_document(
    title: str | None = None,
    include_toc: bool = False,
    logo_path: str | Path | None = None,
    author: str | None = None,
    date: str | None = None,
) -> Document:
    """
    새 Document 객체를 생성합니다.

    권장: 새로운 페이지 구조를 사용하려면 이 함수 대신
    create_document() + build_cover_page() + build_toc_page() + start_body_content()
    조합을 사용하세요.

    Args:
        title: 문서 제목 (None이면 표지 없음)
        include_toc: 목차 포함 여부 (True면 build_toc_page() 자동 호출)
        logo_path: 로고 이미지 경로 (None이면 기본 로고 사용, False면 로고 없음)
        author: 작성자 (None이면 플레이스홀더 표시)
        date: 작성일 (None이면 플레이스홀더 표시)

    Returns:
        초기화된 Document 객체
    """
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    _set_style_korean_font(style)
    style.font.size = DEFAULT_FONT_SIZE
    style.font.color.rgb = COLOR_TEXT_GRAY
    style.paragraph_format.line_spacing = LINE_SPACING

    # Heading 스타일에 브랜드 컬러 적용 및 기본 테두리 제거
    for level, config in HEADING_STYLES.items():
        heading_style_name = f'Heading {level}'
        if heading_style_name in doc.styles:
            h_style = doc.styles[heading_style_name]
            _set_style_korean_font(h_style)
            h_style.font.size = config["size"]
            h_style.font.bold = config["bold"]
            h_style.font.color.rgb = config["color"]

            # 기본 템플릿의 왼쪽 테두리(세로선) 제거
            pPr = h_style._element.get_or_add_pPr()
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                pPr.remove(pBdr)

    # 표지 페이지 생성
    if title:
        build_cover_page(doc, title, logo_path, author, date)

    # 목차 페이지 생성
    if include_toc:
        build_toc_page(doc)

    # 본문 시작 (표지나 목차가 있으면 페이지 브레이크)
    if title or include_toc:
        start_body_content(doc)

    return doc


def build_document(
    sections: list[dict],
    section_results: list[tuple[dict, Route, list[dict]]],
    title: str = "Grounded Report",
    output_path: str = "output.docx",
    logo_path: str | Path | None = None,
    author: str | None = None,
    date: str | None = None,
) -> None:
    """
    섹션 계획과 결과를 받아 문서를 조립합니다.

    Args:
        sections: 섹션 계획 리스트
        section_results: 섹션별 결과
        title: 문서 제목
        output_path: 출력 파일 경로
        logo_path: 로고 이미지 경로 (None이면 기본 로고 사용)
        author: 작성자 (None이면 플레이스홀더 표시)
        date: 작성일 (None이면 플레이스홀더 표시)
    """
    include_toc = len(sections) >= TOC_THRESHOLD
    doc = create_document(
        title=title,
        include_toc=include_toc,
        logo_path=logo_path,
        author=author,
        date=date,
    )

    for section, (llm_result, decision, sources) in zip(sections, section_results):
        add_section(doc, section["title"], llm_result, decision, sources)

    save_document(doc, output_path)


def save_document(doc: Document, filepath: str) -> None:
    """
    Document를 파일로 저장합니다.

    Args:
        doc: Document 객체
        filepath: 저장 경로
    """
    # 디렉토리가 없으면 생성
    output_dir = Path(filepath).parent
    if output_dir and not output_dir.exists():
        output_dir.mkdir(parents=True, exist_ok=True)

    doc.save(filepath)


# =============================================================================
# 유틸리티 함수
# =============================================================================
def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
) -> Table:
    """
    표를 문서에 추가합니다.

    Args:
        doc: python-docx Document 객체
        headers: 헤더 열 리스트
        rows: 데이터 행 리스트

    Returns:
        생성된 Table 객체
    """
    return _add_table_with_header(doc, headers, rows)
