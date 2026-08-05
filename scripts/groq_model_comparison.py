"""
Groq 모델 비교 테스트 스크립트

4개 모델로 동일한 보고서를 생성하여 실행 시간 및 품질 비교
"""

import json
import os
import sys
import time
from datetime import datetime
from typing import Callable

from dotenv import load_dotenv
from groq import Groq

# 상위 디렉토리 추가
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pipeline.rag_tool import search_internal_knowledge
from pipeline.templates import TEMPLATES

load_dotenv()

# 테스트할 모델 목록
MODELS = [
    "llama-3.1-8b-instant",
    "llama-3.3-70b-versatile",
    "openai/gpt-oss-20b",
    "openai/gpt-oss-120b",
]

# Groq 클라이언트
client = Groq(api_key=os.getenv("GROQ_API_KEY"))


SYSTEM_PROMPT = """당신은 보고서 작성을 돕는 AI 어시스턴트입니다.

## 언어 규칙 (절대 준수)
- "answer" 필드는 반드시 한국어로만 작성하세요.
- 중국어(中文), 일본어 등 다른 언어로 절대 전환하지 마세요.
- 영어 기술 용어(AWS, API, Bedrock, LangGraph 등)는 원문 그대로 사용 가능합니다.

중요한 규칙:
1. 반드시 제공된 검색 결과(source) 문서만을 근거로 답변하세요.
2. 검색 결과에 없는 내용은 절대 지어내지 마세요.
3. 근거가 부족하면 "제공된 자료에서 관련 내용을 찾을 수 없습니다"라고 답하세요.

응답 형식:
반드시 아래 JSON 형식으로만 응답하세요.

{
    "answer": "검색 결과를 바탕으로 작성한 답변 내용",
    "source_type": "internal 또는 web",
    "source_count": 참조한 출처 개수,
    "source_relevance": "high/medium/low",
    "has_fabrication_risk": true 또는 false
}"""


SYSTEM_PROMPT_NONE = """당신은 보고서 작성을 돕는 AI 어시스턴트입니다.

## 언어 규칙 (절대 준수)
- "answer" 필드는 반드시 한국어로만 작성하세요.
- 중국어(中文), 일본어 등 다른 언어로 절대 전환하지 마세요.

이 섹션은 서론, 결론 등 검색 결과 없이 작성하는 일반적인 섹션입니다.
문서의 맥락에 맞는 자연스럽고 전문적인 내용을 작성하세요.

응답 형식:
반드시 아래 JSON 형식으로만 응답하세요.

{
    "answer": "작성한 내용",
    "source_type": "none",
    "source_count": 0,
    "source_relevance": "low",
    "has_fabrication_risk": false
}"""


def call_groq(model: str, messages: list[dict]) -> tuple[str, float]:
    """
    Groq API 호출

    Returns:
        (응답 텍스트, 응답 시간 초)
    """
    start = time.time()

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.7,
        max_tokens=2048,
        response_format={"type": "json_object"},
    )

    elapsed = time.time() - start
    content = response.choices[0].message.content
    return content, elapsed


def format_sources(sources: list[dict]) -> str:
    """검색 결과를 LLM에 전달할 형식으로 변환"""
    if not sources:
        return "검색 결과 없음"

    formatted_parts = []
    for i, source in enumerate(sources, 1):
        title = source.get("source_title", f"문서 {i}")
        content = source.get("content", "")
        score = source.get("score", 0.0)
        part = f"[출처 {i}] {title} (관련도: {score:.2f})\n{content}"
        formatted_parts.append(part)

    return "\n\n".join(formatted_parts)


def parse_response(raw: str, source_type: str, source_count: int) -> dict:
    """LLM 응답 파싱"""
    try:
        result = json.loads(raw)
        return {
            "answer": result.get("answer", ""),
            "source_type": result.get("source_type", source_type),
            "source_count": result.get("source_count", source_count),
            "source_relevance": result.get("source_relevance", "low"),
            "has_fabrication_risk": result.get("has_fabrication_risk", False),
        }
    except (json.JSONDecodeError, TypeError):
        return {
            "answer": raw,
            "source_type": source_type,
            "source_count": source_count,
            "source_relevance": "low",
            "has_fabrication_risk": True,
        }


def generate_section(model: str, section_title: str, section_query: str, source_type: str) -> dict:
    """섹션 생성"""
    if source_type == "none":
        messages = [
            {"role": "system", "content": SYSTEM_PROMPT_NONE},
            {"role": "user", "content": f"섹션 요청: {section_query}\n\n위 요청에 맞는 내용을 작성해주세요."},
        ]
        raw, elapsed = call_groq(model, messages)
        result = parse_response(raw, "none", 0)
        result["elapsed"] = elapsed
        return result

    # RAG 검색
    search_results = search_internal_knowledge(section_query, top_k=3)
    sources = search_results["results"]

    formatted_sources = format_sources(sources)
    user_message = f"""질문: {section_query}

검색 결과:
{formatted_sources}

위 검색 결과만을 근거로 질문에 답변해주세요.
source_type 필드는 반드시 "{source_type}"을 사용하세요."""

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": user_message},
    ]

    raw, elapsed = call_groq(model, messages)
    result = parse_response(raw, source_type, len(sources))
    result["elapsed"] = elapsed
    return result


def run_test_for_model(model: str) -> dict:
    """단일 모델로 전체 테스트 실행"""
    print(f"\n{'='*60}")
    print(f"  Model: {model}")
    print(f"{'='*60}")

    # 제안서 템플릿 사용
    template = TEMPLATES["제안서"]
    sections = template["sections"]

    # 섹션별 source_type 매핑 (간단한 규칙 기반)
    source_type_map = {
        "서론": "none",
        "현황 분석": "internal",
        "제안 내용": "internal",
        "기대 효과": "internal",
        "실행 계획": "internal",
        "결론": "none",
    }

    results = {
        "model": model,
        "sections": [],
        "total_time": 0,
        "errors": [],
    }

    total_start = time.time()

    for section in sections:
        title = section["title"]
        query = section["query"]
        source_type = source_type_map.get(title, "internal")

        print(f"\n[{title}] 생성 중...", end=" ", flush=True)

        try:
            section_result = generate_section(model, title, query, source_type)
            section_result["title"] = title
            section_result["query"] = query
            results["sections"].append(section_result)
            print(f"완료 ({section_result['elapsed']:.2f}s)")
        except Exception as e:
            error_msg = f"{title}: {str(e)}"
            results["errors"].append(error_msg)
            print(f"실패 - {e}")

    results["total_time"] = time.time() - total_start

    return results


def save_result_as_docx(result: dict, output_path: str):
    """결과를 docx로 저장"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed, skipping docx generation")
        return

    doc = Document()

    # 제목
    title = doc.add_heading(f"Groq 모델 테스트 결과: {result['model']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 메타 정보
    doc.add_paragraph(f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"총 소요 시간: {result['total_time']:.2f}초")
    doc.add_paragraph()

    # 섹션별 내용
    for section in result["sections"]:
        doc.add_heading(section["title"], level=1)

        # 메타 정보
        meta = doc.add_paragraph()
        meta.add_run(f"[소요시간: {section['elapsed']:.2f}s | ")
        meta.add_run(f"출처유형: {section['source_type']} | ")
        meta.add_run(f"관련성: {section['source_relevance']}]")

        # 본문
        doc.add_paragraph(section["answer"])
        doc.add_paragraph()

    # 에러가 있으면 추가
    if result["errors"]:
        doc.add_heading("오류 목록", level=1)
        for error in result["errors"]:
            doc.add_paragraph(f"- {error}")

    doc.save(output_path)


def print_comparison_table(all_results: list[dict]):
    """비교 테이블 출력"""
    print("\n" + "="*80)
    print("  실행 시간 비교")
    print("="*80)

    # 헤더
    print(f"{'모델':<35} {'총 시간':>10} {'섹션 평균':>10} {'에러':>6}")
    print("-"*80)

    for result in all_results:
        model = result["model"]
        total = result["total_time"]
        sections = result["sections"]
        avg = sum(s["elapsed"] for s in sections) / len(sections) if sections else 0
        errors = len(result["errors"])

        print(f"{model:<35} {total:>8.2f}s {avg:>8.2f}s {errors:>6}")

    print("-"*80)

    # 품질 비교 (샘플 섹션)
    print("\n" + "="*80)
    print("  품질 비교 (현황 분석 섹션)")
    print("="*80)

    for result in all_results:
        model = result["model"]
        # 현황 분석 섹션 찾기
        analysis_section = None
        for s in result["sections"]:
            if s["title"] == "현황 분석":
                analysis_section = s
                break

        if analysis_section:
            print(f"\n[{model}]")
            print(f"  소요시간: {analysis_section['elapsed']:.2f}s")
            print(f"  관련성: {analysis_section['source_relevance']}")
            print(f"  허위생성위험: {analysis_section['has_fabrication_risk']}")
            answer = analysis_section["answer"]
            # 첫 200자만 출력
            preview = answer[:200] + "..." if len(answer) > 200 else answer
            print(f"  내용 미리보기: {preview}")


def main():
    print("="*60)
    print("  Groq 모델 비교 테스트")
    print("  테스트 대상: 4개 모델로 제안서 생성")
    print("="*60)

    all_results = []
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    for model in MODELS:
        try:
            result = run_test_for_model(model)
            all_results.append(result)

            # 각 모델별 docx 저장
            safe_model_name = model.replace("/", "_")
            docx_path = os.path.join(output_dir, f"groq_test_{safe_model_name}_{timestamp}.docx")
            save_result_as_docx(result, docx_path)
            print(f"\n  저장됨: {docx_path}")

        except Exception as e:
            print(f"\n[ERROR] {model}: {e}")
            all_results.append({
                "model": model,
                "sections": [],
                "total_time": 0,
                "errors": [str(e)],
            })

    # 비교 테이블 출력
    print_comparison_table(all_results)

    # 결과 요약 JSON 저장
    summary_path = os.path.join(output_dir, f"groq_test_summary_{timestamp}.json")
    summary = []
    for r in all_results:
        summary.append({
            "model": r["model"],
            "total_time": r["total_time"],
            "section_count": len(r["sections"]),
            "avg_section_time": sum(s["elapsed"] for s in r["sections"]) / len(r["sections"]) if r["sections"] else 0,
            "errors": len(r["errors"]),
        })

    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    print(f"\n\n결과 요약 저장됨: {summary_path}")

    # 생성된 파일 목록 출력
    print("\n" + "="*60)
    print("  생성된 파일들")
    print("="*60)
    for model in MODELS:
        safe_model_name = model.replace("/", "_")
        print(f"  - groq_test_{safe_model_name}_{timestamp}.docx")
    print(f"  - groq_test_summary_{timestamp}.json")


if __name__ == "__main__":
    main()
